"""
Unified file parser for the Portal Context Generator Module.

Supports: CSV, Excel (.xlsx/.xls), Markdown (.md), HTML (.html/.htm),
PDF (.pdf), Word (.docx), and plain text (.txt).

Each file is parsed into a ParsedDocument dataclass containing the extracted
text content, ready to be fed to the LLM for analysis.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pandas as pd
from bs4 import BeautifulSoup

from utils.logger import setup_logger

logger = setup_logger("file_parser")


@dataclass
class ParsedDocument:
    """Represents a parsed document with extracted text content."""
    filename: str
    content: str
    file_type: str
    char_count: int = 0
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        self.char_count = len(self.content)


# ---------------------------------------------------------------------------
# Individual parsers
# ---------------------------------------------------------------------------

def _parse_csv(file_bytes: bytes, filename: str) -> str:
    """Parse CSV file into a markdown table representation."""
    try:
        df = pd.read_csv(io.BytesIO(file_bytes))
        # Limit preview for very large files
        if len(df) > 500:
            content = f"**CSV file with {len(df)} rows and {len(df.columns)} columns.**\n\n"
            content += "**Columns:** " + ", ".join(df.columns.tolist()) + "\n\n"
            content += "**First 100 rows:**\n\n"
            content += df.head(100).to_markdown(index=False)
            content += f"\n\n... and {len(df) - 100} more rows."
        else:
            content = df.to_markdown(index=False)
        return content
    except Exception as e:
        return f"[Error parsing CSV '{filename}': {str(e)}]"


def _parse_excel(file_bytes: bytes, filename: str) -> str:
    """Parse Excel file (multi-sheet support) into markdown tables."""
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes))
        sheets = xls.sheet_names
        parts = [f"**Excel file with {len(sheets)} sheet(s): {', '.join(sheets)}**\n"]

        for sheet_name in sheets:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            parts.append(f"\n### Sheet: {sheet_name}\n")
            if len(df) > 500:
                parts.append(f"*{len(df)} rows, {len(df.columns)} columns. Showing first 100.*\n\n")
                parts.append(df.head(100).to_markdown(index=False))
            else:
                parts.append(df.to_markdown(index=False))

        return "\n".join(parts)
    except Exception as e:
        return f"[Error parsing Excel '{filename}': {str(e)}]"


def _parse_markdown(file_bytes: bytes, filename: str) -> str:
    """Parse Markdown file — direct text read."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Error parsing Markdown '{filename}': {str(e)}]"


def _parse_html(file_bytes: bytes, filename: str) -> str:
    """Parse HTML file — extract text content with structure hints."""
    try:
        soup = BeautifulSoup(file_bytes, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "meta", "link"]):
            element.decompose()

        # Extract title if present
        title = soup.title.string.strip() if soup.title and soup.title.string else ""

        # Get text with some structure preservation
        text = soup.get_text(separator="\n", strip=True)

        if title:
            return f"**Page Title: {title}**\n\n{text}"
        return text
    except Exception as e:
        return f"[Error parsing HTML '{filename}': {str(e)}]"


def _parse_pdf(file_bytes: bytes, filename: str) -> str:
    """Parse PDF file — extract text from all pages."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Page {i} ---\n{text.strip()}")

        if not pages:
            return f"[PDF '{filename}' contained no extractable text — may be image-based]"

        return f"**PDF with {len(reader.pages)} pages:**\n\n" + "\n\n".join(pages)
    except Exception as e:
        return f"[Error parsing PDF '{filename}': {str(e)}]"


def _parse_docx(file_bytes: bytes, filename: str) -> str:
    """Parse Word (.docx) file — extract paragraphs and tables."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading styles
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading", "").strip()
                    try:
                        level = int(level)
                    except ValueError:
                        level = 2
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)

        # Extract tables
        for i, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                parts.append(f"\n**Table {i}:**")
                # Create markdown table
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                parts.append(header)
                parts.append(separator)
                for row in rows[1:]:
                    parts.append("| " + " | ".join(row) + " |")

        return "\n\n".join(parts) if parts else f"[DOCX '{filename}' contained no extractable content]"
    except Exception as e:
        return f"[Error parsing DOCX '{filename}': {str(e)}]"


def _parse_text(file_bytes: bytes, filename: str) -> str:
    """Parse plain text file — direct read."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Error parsing text file '{filename}': {str(e)}]"


# ---------------------------------------------------------------------------
# File type detection & unified parser
# ---------------------------------------------------------------------------

# Map of file extensions to their parser functions
PARSER_MAP = {
    ".csv": ("csv", _parse_csv),
    ".xlsx": ("excel", _parse_excel),
    ".xls": ("excel", _parse_excel),
    ".md": ("markdown", _parse_markdown),
    ".markdown": ("markdown", _parse_markdown),
    ".html": ("html", _parse_html),
    ".htm": ("html", _parse_html),
    ".pdf": ("pdf", _parse_pdf),
    ".docx": ("docx", _parse_docx),
    ".txt": ("text", _parse_text),
    ".text": ("text", _parse_text),
    ".log": ("text", _parse_text),
    ".json": ("text", _parse_text),
    ".xml": ("html", _parse_html),  # Treat XML like HTML for text extraction
}

SUPPORTED_EXTENSIONS = list(PARSER_MAP.keys())


def parse_file(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Parse a file and return a ParsedDocument.

    Args:
        file_bytes: Raw file content as bytes.
        filename: Original filename (used for type detection and metadata).

    Returns:
        ParsedDocument with extracted text content.

    Raises:
        ValueError: If the file type is not supported.
    """
    ext = Path(filename).suffix.lower()

    if ext not in PARSER_MAP:
        raise ValueError(
            f"Unsupported file type '{ext}' for file '{filename}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    file_type, parser_fn = PARSER_MAP[ext]
    logger.debug(f"Parsing file '{filename}' as {file_type}")
    content = parser_fn(file_bytes, filename)

    return ParsedDocument(
        filename=filename,
        content=content,
        file_type=file_type,
        metadata={"extension": ext},
    )


def parse_multiple_files(
    files: list[tuple[bytes, str]],
) -> list[ParsedDocument]:
    """
    Parse multiple files and return a list of ParsedDocuments.

    Args:
        files: List of (file_bytes, filename) tuples.

    Returns:
        List of ParsedDocument objects. Failed files are included with error messages.
    """
    documents = []
    for file_bytes, filename in files:
        try:
            logger.info(f"Attempting to parse: {filename}")
            doc = parse_file(file_bytes, filename)
            logger.info(f"Successfully parsed {filename} ({doc.char_count} chars)")
            documents.append(doc)
        except ValueError as e:
            # Include unsupported files with an error message so the user sees feedback
            documents.append(ParsedDocument(
                filename=filename,
                content=f"[UNSUPPORTED: {str(e)}]",
                file_type="unsupported",
                metadata={"error": str(e)},
            ))
    return documents


def get_supported_formats_display() -> str:
    """Return a human-readable string of supported file formats."""
    formats = {
        "CSV": [".csv"],
        "Excel": [".xlsx", ".xls"],
        "Markdown": [".md", ".markdown"],
        "HTML": [".html", ".htm"],
        "PDF": [".pdf"],
        "Word": [".docx"],
        "Text": [".txt", ".text", ".log"],
        "JSON": [".json"],
        "XML": [".xml"],
    }
    parts = [f"**{name}** ({', '.join(exts)})" for name, exts in formats.items()]
    return " · ".join(parts)
